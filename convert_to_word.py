#!/usr/bin/env python3
"""
Improved Markdown → Word converter
- Handles bullets, numbered lists
- Inline bold/italic/code
- Code blocks and file-tree formatting
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def set_doc_defaults(doc):
    # Normal font
    try:
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
    except Exception:
        pass
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

def add_code_block(doc, code_lines):
    # Add a monospace paragraph with smaller font and a light gray shading (best-effort)
    para = doc.add_paragraph()
    run = para.add_run('\n'.join(code_lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    para.style = 'No Spacing'
    # (Advanced shading could be added via OXML if desired.)

def is_tree_line(line):
    return any(ch in line for ch in ('├──', '└──', '│', '─'))

_inline_pattern = re.compile(
    r'(`[^`]+`)|'        # inline code (`code`)
    r'(\*\*[^*]+\*\*)|'  # bold (**bold**)
    r'(\*[^*]+\*)|'      # italic (*italic*)
    r'(\[([^\]]+)\]\(([^)]+)\))'  # link [text](url) -- rendered as text
)

def add_paragraph_with_format(doc, text, style=None, monospace=False):
    """
    Add a paragraph but parse inline markdown tokens for bold/italic/code/links.
    monospace=True forces Courier New for normal text too.
    """
    if style:
        para = doc.add_paragraph(style=style)
    else:
        para = doc.add_paragraph()
    pos = 0
    for m in _inline_pattern.finditer(text):
        start, end = m.span()
        # text before match
        if start > pos:
            segment = text[pos:start]
            r = para.add_run(segment)
            if monospace:
                r.font.name = 'Courier New'
                r.font.size = Pt(10)
        matched = m.group(0)
        # inline code
        if m.group(1):
            inner = matched.strip('`')
            r = para.add_run(inner)
            r.font.name = 'Courier New'
            r.font.size = Pt(10)
        # bold
        elif m.group(2):
            inner = matched[2:-2]
            r = para.add_run(inner)
            r.bold = True
        # italic
        elif m.group(3):
            inner = matched[1:-1]
            r = para.add_run(inner)
            r.italic = True
        # link [text](url) -> show text (url) (you can make real hyperlink if needed)
        elif m.group(4):
            textpart = m.group(5)  # visible text
            urlpart = m.group(6)
            r = para.add_run(f"{textpart} ({urlpart})")
            r.underline = True
        pos = end
    # trailing text
    if pos < len(text):
        tail = text[pos:]
        r = para.add_run(tail)
        if monospace:
            r.font.name = 'Courier New'
            r.font.size = Pt(10)

def convert_md_to_word(md_file, word_file):
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = Document()
    set_doc_defaults(doc)
    title = doc.add_heading('Trading Platform API Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    i = 0
    in_code_block = False
    code_acc = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip('\n')
        stripped = line.strip()

        # Toggle code fence
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_acc = []
            else:
                # close code block
                in_code_block = False
                add_code_block(doc, code_acc)
                code_acc = []
            i += 1
            continue

        if in_code_block:
            code_acc.append(line)
            i += 1
            continue

        # skip empty lines (but allow vertical spacing)
        if stripped == '':
            i += 1
            continue

        # Headers
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:].strip(), level=4)
            i += 1
            continue

        # File tree lines: keep monospace and do not try to interpret '*'
        if is_tree_line(line):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = 'Courier New'
            r.font.size = Pt(10)
            i += 1
            continue

        # Ordered list
        mnum = re.match(r'^\s*(\d+)\.\s+(.*)', line)
        if mnum:
            content = mnum.group(2).strip()
            # remove surrounding backticks from filenames (e.g. `file.py`) for clean text
            content = re.sub(r'`([^`]+)`', r'\1', content)
            add_paragraph_with_format(doc, content, style='List Number')
            i += 1
            continue

        # Unordered list (handles -, *, •)
        mlist = re.match(r'^\s*([\-*\u2022])\s+(.*)', line)
        if mlist:
            content = mlist.group(2).strip()
            content = re.sub(r'`([^`]+)`', r'\1', content)  # strip backticks around filenames
            add_paragraph_with_format(doc, content, style='List Bullet')
            i += 1
            continue

        # Inline code-only line (like `module.py`: desc) — strip backticks and format
        if '`' in line and ':' in line and line.count('`') >= 2:
            # split like `file.py`: description
            line2 = re.sub(r'`([^`]+)`', r'\1', line)
            add_paragraph_with_format(doc, line2)
            i += 1
            continue

        # Status indicators or lines with emoji - keep as bullet for readability
        if any(sym in line for sym in ('✅', '❌', '🔄')):
            add_paragraph_with_format(doc, line.strip(), style='List Bullet')
            i += 1
            continue

        # Lines with inline backticks (inline code) or bold/italic
        if '`' in line or '**' in line or '*' in line or '[' in line:
            # remove stray leading/trailing asterisks that are not markup (e.g., separators)
            cleaned = line.strip()
            # If the whole line is like "* something *" (surrounded by stars), strip only the leading list star (handled above)
            cleaned = re.sub(r'^\*+\s*', '', cleaned)
            cleaned = re.sub(r'\s*\*+$', '', cleaned)
            add_paragraph_with_format(doc, cleaned)
            i += 1
            continue

        # Fallback: plain paragraph (also strip unnecessary backticks around single words)
        cleaned = re.sub(r'`([^`]+)`', r'\1', line)
        # Remove leftover leading/trailing stars that are typical artifacts
        cleaned = cleaned.strip()
        cleaned = re.sub(r'^\*+\s*', '', cleaned)
        cleaned = re.sub(r'\s*\*+$', '', cleaned)
        if cleaned:
            add_paragraph_with_format(doc, cleaned)
        i += 1

    # Save
    doc.save(word_file)
    print(f"✅ Converted {md_file} -> {word_file}")

def main():
    try:
        convert_md_to_word('pretrained_model_analysis.md', 'Why not use pretrained models.docx')
        print("🎉 Done.")
    except FileNotFoundError:
        print("❌ API_DOCUMENTATION.md not found")
    except Exception as e:
        print(f"❌ Error: {e}")

if __name__ == '__main__':
    main()
